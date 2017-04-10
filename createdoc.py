from docx import Document
from docx.shared import Inches
import barcode
from barcode.writer import ImageWriter
label = "704041"
CODE39 = barcode.get_barcode_class('code39')
CODE39.default_writer_options['module_height'] = 3.0
CODE39.default_writer_options['module_width'] = 0.01
CODE39.default_writer_options['text_distance'] = 0.3
CODE39.default_writer_options['font_size'] = 16
#ean = barcode.get_barcode('code39', label, writer=ImageWriter())
#ean.default_writer_options['module_height'] = 3.0
#ean.default_writer_options['add_checksum'] = False
#ean.default_writer_options[''] = 0.2
#ean.calculate_size(dpb)
code39 = CODE39(label, writer=ImageWriter(),add_checksum=False)
pngfile='C:\\data\\' + label
filename = code39.save(pngfile)
#filename = ean.save(pngfile)

document = Document()
sections = document.sections
for section in sections:
	section.top_margin = Inches(0.1)
	section.bottom_margin = Inches(0.2)
	section.left_margin = Inches(0.25)
	section.right_margin = Inches(0.3)

p = document.add_paragraph()
r = p.add_run()

r.add_picture(filename,width=Inches(1.50) )
r.add_picture(filename,width=Inches(1.50) )
r.add_picture(filename,width=Inches(1.50) )
r.add_picture(filename,width=Inches(1.50) )
r.add_picture(filename,width=Inches(1.70) )
r.add_picture(filename,width=Inches(1.70) )
r.add_picture(filename,width=Inches(1.70) )
#r.add_picture(filename)
#r.add_picture(filename)
#r.add_picture(filename)
#r.add_picture(filename)
#r.add_picture(filename)
#r.add_picture(filename)
#r.add_picture(filename)

document.add_picture(filename, width=Inches(1.70))
document.add_picture(filename, width=Inches(1.70))
#r.add_text(' do you like it?')

document.save('demo.docx')